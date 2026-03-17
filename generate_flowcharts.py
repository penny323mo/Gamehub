import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import matplotlib.patheffects as pe
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# Use Noto Sans CJK font
plt.rcParams['font.family'] = 'Noto Sans CJK JP'
plt.rcParams['axes.unicode_minus'] = False

# ─────────────────────────────────────────────
# Drawing helpers
# ─────────────────────────────────────────────

def draw_rounded_rect(ax, x, y, w, h, color, text, fontsize=8.5, text_color='black', radius=0.3):
    """Draw a rounded rectangle with centred text."""
    box = FancyBboxPatch((x - w/2, y - h/2), w, h,
                         boxstyle=f"round,pad=0.05,rounding_size={radius}",
                         facecolor=color, edgecolor='#444444', linewidth=1.2, zorder=3)
    ax.add_patch(box)
    ax.text(x, y, text, ha='center', va='center', fontsize=fontsize,
            color=text_color, wrap=True, zorder=4,
            multialignment='center')

def draw_diamond(ax, x, y, w, h, color, text, fontsize=8.5):
    """Draw a diamond shape with centred text."""
    diamond = plt.Polygon([[x, y+h/2], [x+w/2, y], [x, y-h/2], [x-w/2, y]],
                           facecolor=color, edgecolor='#444444', linewidth=1.2, zorder=3)
    ax.add_patch(diamond)
    ax.text(x, y, text, ha='center', va='center', fontsize=fontsize, zorder=4,
            multialignment='center')

def draw_stadium(ax, x, y, w, h, color, text, fontsize=9, text_color='white'):
    """Draw a stadium/oval shape (start/end)."""
    ellipse = mpatches.Ellipse((x, y), w, h,
                               facecolor=color, edgecolor='#444444', linewidth=1.5, zorder=3)
    ax.add_patch(ellipse)
    ax.text(x, y, text, ha='center', va='center', fontsize=fontsize,
            color=text_color, fontweight='bold', zorder=4)

def arrow(ax, x1, y1, x2, y2, label='', color='#444444', style='->', dashed=False):
    """Draw an arrow between two points."""
    ls = '--' if dashed else '-'
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle='->', color=color, lw=1.3,
                                linestyle=ls),
                zorder=2)
    if label:
        mx, my = (x1+x2)/2, (y1+y2)/2
        ax.text(mx+0.08, my, label, fontsize=8, color='#666666', zorder=5)

def arrow_h(ax, x1, y1, x2, y2, label='', dashed=False):
    """Horizontal then vertical arrow (L-shape)."""
    ls = '--' if dashed else '-'
    # draw polyline manually
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle='->', color='#444444', lw=1.3,
                                linestyle=ls,
                                connectionstyle='angle,angleA=0,angleB=90,rad=5'),
                zorder=2)
    if label:
        ax.text((x1+x2)/2, (y1+y2)/2+0.05, label, fontsize=8, color='#666666', zorder=5)


# ─────────────────────────────────────────────
# Colour palette
# ─────────────────────────────────────────────
C_START  = '#2E86AB'   # blue  – start/end
C_ACTION = '#F5F5F5'   # light grey – action box
C_ADMIN  = '#E8F4F8'   # pale blue – admin step
C_DIAMOND= '#FFF3CD'   # pale yellow – decision
C_SIDE   = '#E8F5E9'   # pale green – side branch


# ═══════════════════════════════════════════════════════════
# FLOWCHART 1 – 各類申請證明書
# ═══════════════════════════════════════════════════════════
def make_chart1():
    fig, ax = plt.subplots(figsize=(11, 20))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 20)
    ax.axis('off')
    fig.patch.set_facecolor('white')

    # ── Node positions (x, y) ──────────────────────────────
    #   Main spine: centre x=5
    cx = 5.5

    yA  = 19.2   # 開始
    yB  = 18.0   # 申請人自行申請
    yC  = 16.5   # 是否需要上級確認？
    yD  = 15.0   # 上級提供意見  (right branch x=8.5)
    yE  = 13.5   # 行政申請登記
    yF  = 12.0   # 是否需要交相關人員檢視？
    yG  = 10.5   # 相關負責人員檢視  (right branch x=8.5)
    yH  =  9.0   # 行政文件：處長及秘書
    yI  =  7.8   # 行政文件：廳長及秘書
    yJ  =  6.6   # 領導層審批
    yK  =  5.4   # 行政申請－人事及行政處
    yL  =  4.2   # 行政文件：處長及秘書
    yM  =  3.0   # 行政文件：廳長及秘書
    yN  =  1.8   # 歸檔
    yO  =  0.8   # 結束

    # Side branch
    xSide = 9.2
    yP   =  3.6   # 秘書調閱文件
    yQ   =  2.4   # 按需要分派

    bw, bh = 3.6, 0.65   # box width/height
    dw, dh = 3.8, 0.9    # diamond

    # ── Nodes ──────────────────────────────────────────────
    draw_stadium(ax, cx, yA, 2.0, 0.7, C_START, '開始')
    draw_rounded_rect(ax, cx, yB, bw, bh, C_ACTION, '申請人自行申請')
    draw_diamond(ax, cx, yC, dw, dh, C_DIAMOND, '是否需要上級確認？')

    # right branch: D
    draw_rounded_rect(ax, 8.5, yD, 3.0, bh, C_ACTION, '上級提供意見')

    draw_rounded_rect(ax, cx, yE, bw, bh, C_ADMIN,
                      '行政申請－人事及行政處\n登記申請', fontsize=8)
    draw_diamond(ax, cx, yF, dw+0.2, dh, C_DIAMOND, '是否需要交相關人員檢視？')

    # right branch: G
    draw_rounded_rect(ax, 8.5, yG, 3.0, bh, C_ACTION, '相關負責人員檢視／核對資料', fontsize=8)

    draw_rounded_rect(ax, cx, yH, bw, bh, C_ADMIN,
                      '行政文件：處長及秘書\n人事及行政處', fontsize=8)
    draw_rounded_rect(ax, cx, yI, bw, bh, C_ADMIN,
                      '行政文件：廳長及秘書\n管理及計劃廳', fontsize=8)
    draw_rounded_rect(ax, cx, yJ, bw, bh, C_ACTION, '領導層審批')
    draw_rounded_rect(ax, cx, yK, bw, bh, C_ADMIN,
                      '行政申請－人事及行政處', fontsize=8)
    draw_rounded_rect(ax, cx, yL, bw, bh, C_ADMIN,
                      '行政文件：處長及秘書\n人事及行政處', fontsize=8)
    draw_rounded_rect(ax, cx, yM, bw, bh, C_ADMIN,
                      '行政文件：廳長及秘書\n管理及計劃廳', fontsize=8)
    draw_rounded_rect(ax, cx, yN, bw, bh, C_ACTION, '歸檔')
    draw_stadium(ax, cx, yO, 2.0, 0.7, C_START, '結束')

    # Side branch nodes
    draw_rounded_rect(ax, xSide, yP, 2.8, bh, C_SIDE, '秘書調閱文件及列印', fontsize=8)
    draw_rounded_rect(ax, xSide, yQ, 2.8, bh, C_SIDE, '按需要分派相關同事跟進', fontsize=8)

    # ── Arrows: main spine ────────────────────────────────
    arrow(ax, cx, yA-0.35, cx, yB+bh/2)
    arrow(ax, cx, yB-bh/2, cx, yC+dh/2)

    # Diamond C -> E (否, left/straight down)
    arrow(ax, cx, yC-dh/2, cx, yE+bh/2, label='否')

    # Diamond C -> D (是, right)
    ax.annotate('', xy=(8.5, yD+bh/2), xytext=(cx+dw/2, yC),
                arrowprops=dict(arrowstyle='->', color='#444444', lw=1.3,
                                connectionstyle='angle,angleA=0,angleB=-90,rad=3'), zorder=2)
    ax.text(cx+dw/2+0.1, yC+0.15, '是', fontsize=8, color='#666666')

    # D -> E (right to main)
    ax.annotate('', xy=(cx+bw/2, yE), xytext=(8.5, yD-bh/2),
                arrowprops=dict(arrowstyle='->', color='#444444', lw=1.3,
                                connectionstyle='angle,angleA=0,angleB=90,rad=3'), zorder=2)

    arrow(ax, cx, yE-bh/2, cx, yF+dh/2)

    # Diamond F -> H (不需要)
    arrow(ax, cx, yF-dh/2, cx, yH+bh/2, label='不需要')

    # Diamond F -> G (需要, right)
    ax.annotate('', xy=(8.5, yG+bh/2), xytext=(cx+dw/2+0.1, yF),
                arrowprops=dict(arrowstyle='->', color='#444444', lw=1.3,
                                connectionstyle='angle,angleA=0,angleB=-90,rad=3'), zorder=2)
    ax.text(cx+dw/2+0.2, yF+0.15, '需要', fontsize=8, color='#666666')

    # G -> H
    ax.annotate('', xy=(cx+bw/2, yH), xytext=(8.5, yG-bh/2),
                arrowprops=dict(arrowstyle='->', color='#444444', lw=1.3,
                                connectionstyle='angle,angleA=0,angleB=90,rad=3'), zorder=2)

    arrow(ax, cx, yH-bh/2, cx, yI+bh/2)
    arrow(ax, cx, yI-bh/2, cx, yJ+bh/2)
    arrow(ax, cx, yJ-bh/2, cx, yK+bh/2)
    arrow(ax, cx, yK-bh/2, cx, yL+bh/2)
    arrow(ax, cx, yL-bh/2, cx, yM+bh/2)
    arrow(ax, cx, yM-bh/2, cx, yN+bh/2)
    arrow(ax, cx, yN-bh/2, cx, yO+0.35)

    # Side branch: L -.-> P -.-> Q
    ax.annotate('', xy=(xSide, yP+bh/2), xytext=(cx+bw/2, yL),
                arrowprops=dict(arrowstyle='->', color='#555555', lw=1.1,
                                linestyle='--',
                                connectionstyle='angle,angleA=0,angleB=-90,rad=3'), zorder=2)
    arrow(ax, xSide, yP-bh/2, xSide, yQ+bh/2, dashed=True)

    # Title
    ax.text(cx, 19.8, '流程圖一：各類申請證明書',
            ha='center', va='center', fontsize=13, fontweight='bold', color='#1a1a2e')

    plt.tight_layout(pad=0.5)
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=160, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════════════════════════
# FLOWCHART 2 – 出勤聲明書
# ═══════════════════════════════════════════════════════════
def make_chart2():
    fig, ax = plt.subplots(figsize=(10, 18))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 18)
    ax.axis('off')
    fig.patch.set_facecolor('white')

    cx = 5.0
    bw, bh = 3.6, 0.65

    yA = 17.2
    yB = 16.0
    yC = 14.8
    yD = 13.6
    yE = 12.2
    yF = 10.8
    yG =  9.5
    yH =  8.2
    yI =  7.0
    yJ =  5.8
    yK =  4.6
    yL =  3.4
    yM =  2.2
    yN =  1.0

    draw_stadium(ax, cx, yA, 2.0, 0.7, C_START, '開始')
    draw_rounded_rect(ax, cx, yB, bw, bh, C_ACTION, '申請人提出出勤聲明申請')
    draw_rounded_rect(ax, cx, yC, bw, bh, C_ACTION, '主管審閱／確認')
    draw_rounded_rect(ax, cx, yD, bw, bh, C_ACTION, '所屬處／廳審閱')
    draw_rounded_rect(ax, cx, yE, bw, bh, C_ADMIN,
                      '行政申請－人事及行政處\n登記申請', fontsize=8)
    draw_rounded_rect(ax, cx, yF, bw, bh, C_ACTION,
                      '專責處理出勤之人員\n核實資料並上載所需附件', fontsize=8)
    draw_rounded_rect(ax, cx, yG, bw, bh, C_ADMIN,
                      '行政文件：處長及秘書\n人事及行政處', fontsize=8)
    draw_rounded_rect(ax, cx, yH, bw, bh, C_ADMIN,
                      '行政文件：廳長及秘書\n管理及計劃廳', fontsize=8)
    draw_rounded_rect(ax, cx, yI, bw, bh, C_ACTION, '領導層審批')
    draw_rounded_rect(ax, cx, yJ, bw, bh, C_ADMIN,
                      '行政申請－人事及行政處', fontsize=8)
    draw_rounded_rect(ax, cx, yK, bw, bh, C_ADMIN,
                      '行政文件：處長及秘書\n人事及行政處', fontsize=8)
    draw_rounded_rect(ax, cx, yL, bw, bh, C_ADMIN,
                      '行政文件：廳長及秘書\n管理及計劃廳', fontsize=8)
    draw_rounded_rect(ax, cx, yM, bw, bh, C_ACTION, '歸檔')
    draw_stadium(ax, cx, yN, 2.0, 0.7, C_START, '結束')

    for ya, yb in [(yA, yB), (yB, yC), (yC, yD), (yD, yE),
                   (yE, yF), (yF, yG), (yG, yH), (yH, yI),
                   (yI, yJ), (yJ, yK), (yK, yL), (yL, yM)]:
        top = ya - (0.35 if ya in (yA,) else bh/2)
        bot = yb + (0.35 if yb in (yN,) else bh/2)
        arrow(ax, cx, top, cx, bot)
    arrow(ax, cx, yM-bh/2, cx, yN+0.35)

    ax.text(cx, 17.8, '流程圖二：出勤聲明書',
            ha='center', va='center', fontsize=13, fontweight='bold', color='#1a1a2e')

    plt.tight_layout(pad=0.5)
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=160, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════════════════════════
# Build Word document
# ═══════════════════════════════════════════════════════════
def build_word():
    doc = Document()

    # Page setup – A4 landscape for chart1, portrait for chart2
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    section = doc.sections[0]
    section.page_width  = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin   = Inches(0.75)
    section.right_margin  = Inches(0.75)
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # ── Chart 1 ──────────────────────────────────────────
    h = doc.add_heading('流程圖一：各類申請證明書', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    buf1 = make_chart1()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf1, width=Inches(7.2))

    doc.add_page_break()

    # ── Chart 2 ──────────────────────────────────────────
    h2 = doc.add_heading('流程圖二：出勤聲明書', level=1)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    buf2 = make_chart2()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run()
    run2.add_picture(buf2, width=Inches(6.5))

    out = '/home/user/Gamehub/flowcharts.docx'
    doc.save(out)
    print(f'Saved: {out}')

build_word()
